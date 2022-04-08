# coding=utf-8

from docx import Document
import pdbc
from file_type_convert import file_type_convert

from config import Config
from utils import Utils

class DocTrans(object):
    def __init__(self) -> None:
        self.p = pdbc.PdbcConnector()
        self.file_convert = file_type_convert()
        self.utils = Utils()

    def docTrans(self, file_path, document_save_name, src, tgt, document_id, user_id):
        doc = Document(file_path)

        id_index = 0
        source_list = []
        source_len = 0
        trans_result = []

        for para in doc.paragraphs:
            source_text = ""
            for run in para.runs:
                source = run.text
                if source.lstrip().rstrip() != '':
                    source_text += source
            if source_text:
                data = {}
                data['id'] = id_index
                data['text'] = source_text
                id_index += 1
                source_list.append(data)
                source_len += len(source_text)

            if source_len > Config.MAX_CHARACTERS:
                result = self.utils.getNmtJson(src, tgt, source_list, user_id)
                #print(result)
                trans_result.extend(result)
                source_len = 0
                source_list = []

        if source_len != 0:
            result = self.utils.getNmtJson(src, tgt, source_list, user_id)
            #print(result)
            trans_result.extend(result)

        for para in doc.paragraphs:
            if para.text.lstrip().rstrip() != '':
                para.text = para.text.replace(para.text, trans_result[0]['trans_text'])
                del (trans_result[0])

        id_index = 0
        source_list = []
        source_len = 0
        trans_result = []

        for t in doc.tables:
            for row in t.rows:
                source_text = ""
                for cell in row.cells:
                    source = cell.text
                    if source.lstrip().rstrip() != '':
                        source_text = source

                    if source_text:
                        data = {}
                        data['id'] = id_index
                        data['text'] = source_text
                        id_index += 1
                        source_list.append(data)
                        source_len += len(source_text)

                    if source_len > Config.MAX_CHARACTERS:
                        result = self.utils.getNmtJson(src, tgt, source_list, user_id)
                        trans_result.extend(result)
                        source_len = 0
                        source_list = []

        if source_len != 0:
            result = self.utils.getNmtJson(src, tgt, source_list, user_id)
            trans_result.extend(result)

        for t in doc.tables:
            for row in t.rows:
                for cell in row.cells:
                    if cell.text.lstrip().rstrip() != '':
                        cell.text = cell.text.replace(cell.text, trans_result[0]['trans_text'])
                        del (trans_result[0])

        try:
            doc.save(Config.DOC_RESULT_FOLDER + document_save_name)
        except Exception as e:
            print(e)
            return False

        self.p.updateDocumentStatus(document_id, '2')
        return True

