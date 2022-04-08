#!/usr/bin/env python 
# -*- coding: utf-8 -*-
# @Time    : 2021/7/7 21:35
# @Author  : Rebekah Jiang
# @File    : file_trans.py
# @Software: PyCharm

import os
from docx import Document
import pdbc
from file_type_convert import file_type_convert
from openpyxl import load_workbook
from cutString import cutString
from clean_f import clean_f
from merge import merge
from config import Config
from utils import Utils
import threading
import math
import time


class fileTrans(object):
    def __init__(self):
        self.p = pdbc.PdbcConnector()
        self.file_convert = file_type_convert()
        self.cutString = cutString()
        self.clean_f = clean_f()
        self.merge = merge()
        self.utils = Utils()

    def docTrans(self, file_path, document_save_name, src, tgt, document_id):
        doc = Document(file_path)

        id_index = 0
        source_list = []
        source_len = 0
        trans_result = []

        for para in doc.paragraphs:
            source_text = ""
            for run in para.runs:
                source = run.text
                if source.lstrip().rstrip() != '':
                    source_text += source
            if source_text:
                data = {}
                data['id'] = id_index
                data['text'] = source_text
                id_index += 1
                source_list.append(data)
                source_len += len(source_text)

            if source_len > Config.MAX_CHARACTERS:
                result = self.utils.getNmtJson(src, tgt, source_list)
                #print(result)
                trans_result.extend(result)
                source_len = 0
                source_list = []

        if source_len != 0:
            result = self.utils.getNmtJson(src, tgt, source_list)
            #print(result)
            trans_result.extend(result)

        for para in doc.paragraphs:
            if para.text.lstrip().rstrip() != '':
                para.text = para.text.replace(para.text, trans_result[0]['trans_text'])
                del (trans_result[0])

        id_index = 0
        source_list = []
        source_len = 0
        trans_result = []

        for t in doc.tables:
            for row in t.rows:
                source_text = ""
                for cell in row.cells:
                    source = cell.text
                    if source.lstrip().rstrip() != '':
                        source_text = source

                    if source_text:
                        data = {}
                        data['id'] = id_index
                        data['text'] = source_text
                        id_index += 1
                        source_list.append(data)
                        source_len += len(source_text)

                    if source_len > Config.MAX_CHARACTERS:
                        result = self.utils.getNmtJson(src, tgt, source_list)
                        trans_result.extend(result)
                        source_len = 0
                        source_list = []

        if source_len != 0:
            result = self.utils.getNmtJson(src, tgt, source_list)
            trans_result.extend(result)

        for t in doc.tables:
            for row in t.rows:
                for cell in row.cells:
                    if cell.text.lstrip().rstrip() != '':
                        cell.text = cell.text.replace(cell.text, trans_result[0]['trans_text'])
                        del (trans_result[0])

        try:
            doc.save(Config.DOC_RESULT_FOLDER + document_save_name)
        except Exception as e:
            print(e)
            return False

        self.p.updateDocumentStatus(document_id, '2')
        return True

    def sheetThread(self, wb, sheet, document_save_name, src, tgt):
        colThreads = []
        for col in sheet.iter_cols():
            aThread = threading.Thread(target=self.excelTransThread,
                                       args=(wb, col, document_save_name, src, tgt))
            colThreads.append(aThread)

        for thread in colThreads:
            thread.start()

        for t in colThreads:
            t.join()

    def excelTransThread(self, wb, col, document_save_name, src, tgt):
        id_index = 0
        source_len = 0
        source_list = []
        trans_list = []
        trans_result = []
        source_content = ""

        for cell in col:
            if cell.value is not None and str(cell.value).isdigit() == False:
                clean_out = self.clean_f.clean(cell.value)
                keep_str = []
                if clean_out.lstrip().rstrip() == '':
                    keep_str.append(cell.value)
                    source_list.append(keep_str)
                    continue
                cut_out = self.cutString.get_cut_word(clean_out)
                if cut_out[0][1] == "blank":
                    cut_out.pop(0)
                merge_out = self.merge.merge(cut_out)
                source_list.append(merge_out)
                source_content = ""
                for each in merge_out:
                    each = list(each)
                    if each[1] == src:
                        source_content = source_content + each[0] + "\n"

                data = {}
                data['id'] = id_index
                data['text'] = source_content
                trans_list.append(data)
                id_index += 1
                source_len += len(source_content)

                if source_len > Config.MAX_CHARACTERS:
                    result = self.utils.getNmtJson(src, tgt, trans_list)
                    trans_result.extend(result)
                    source_len = 0
                    trans_list = []

        if source_len != 0:
            result = self.utils.getNmtJson(src, tgt, trans_list)
            trans_result.extend(result)

        trans_result_list = []
        for each in trans_result:
            trans_result_list.extend(each['trans_text'].split("\n"))

        for cell in col:
            if cell.value is not None and str(cell.value).isdigit() == False:
                value = ""
                for cut_untrans in source_list[0]:
                    if len(source_list[0]) == 1 and isinstance(cut_untrans, str):
                        value = cut_untrans
                        continue
                    elif cut_untrans[1] == src:
                        value += trans_result_list[0]
                        del (trans_result_list[0])
                    else:
                        value += cut_untrans[0]
                cell.value = value
                del (source_list[0])
        wb.save(Config.DOC_RESULT_FOLDER + document_save_name)

    def excelTrans(self, file_path, document_save_name, src, tgt, document_id):
        # p = pdbc.PdbcConnector()
        wb = load_workbook(filename=file_path)
        # ws = wb.active

        sheetThreads = []
        for sheet in wb:
            aThread = threading.Thread(target=self.sheetThread,
                                       args=(wb, sheet, document_save_name, src, tgt))
            sheetThreads.append(aThread)

        for thread in sheetThreads:
            thread.start()

        for t in sheetThreads:
            t.join()
        try:
            wb.save(Config.DOC_RESULT_FOLDER + document_save_name)
        except Exception as e:
            print(e)
            return False

        self.p.updateDocumentStatus(document_id, '2')
        return True

    def txtTrans(self, file_path, document_save_name, src, tgt, document_id, isSplit):
        print("############ translate file " + file_path + " ################")
        with open(file_path, 'r') as fr:
            id_index = 0
            source_list = []
            source_len = 0
            trans_result = []

            for source_text in fr.readlines():
                if source_text.lstrip().rstrip() != '':
                    data = {}
                    data['id'] = id_index
                    data['text'] = source_text
                    id_index += 1
                    source_list.append(data)
                    source_len += len(source_text)
                
                if source_len > Config.MAX_CHARACTERS:
                    result = self.utils.getNmtJson(src, tgt, source_list)
                        # print(result)
                    if not result:
                        time.sleep(5)
                        result = self.utils.getNmtJson(src, tgt, source_list)
                        if not result:
                            result = source_list
                        # try:
                        #     result = self.utils.getNmtJson(src, tgt, source_list)
                        #     # print(result)
                        # except Exception as e:
                        #     print(e)
                        #     # return False
                        #     result = source_list
                    trans_result.extend(result)
                    source_len = 0
                    source_list = []
            
            if source_len != 0:
                result = self.utils.getNmtJson(src, tgt, source_list)
                if not result:
                    time.sleep(5)
                    result = self.utils.getNmtJson(src, tgt, source_list)
                    if not result:
                        result = source_list
                    # try:
                    #     result = self.utils.getNmtJson(src, tgt, source_list)
                    #     # print(result)
                    # except Exception as e:
                    #     print(e)
                    #     # return False
                    #     result = source_list
                trans_result.extend(result)
        
        try:
            save_path = Config.DOC_RESULT_FOLDER + document_save_name
            if isSplit:
                save_path = os.path.join(os.path.split(file_path)[0], 'concat', os.path.split(file_path)[1])
            
            with open(file_path, 'r') as fr, open(save_path, 'w') as fw:
                for source_text in fr:
                    if source_text.lstrip().rstrip() != '':
                        replace_text = trans_result[0]['trans_text'] if 'trans_text' in trans_result[0] else trans_result[0]['text']
                        source_text = source_text.replace(source_text, replace_text)
                        del (trans_result[0])
                    fw.write(source_text)
                # for tgt_text in trans_result:
                #     fw.write(tgt_text['trans_text'])
        except Exception as e:
            print(e)
            return False

        # self.p.updateDocumentStatus(document_id, '2')
        return True

    def fileTransThread(self, document_id, user_id, document_save_name, src_lang, tgt_lang, document_name, document_size):
        file_type = os.path.splitext(document_save_name)[1]

        if file_type == '.doc' or file_type == '.docx':
            path, document_save_name = self.file_convert.doc_to_docx(Config.DOC_UPLOAD_FOLDER, document_save_name)
            self.docTrans(path, document_save_name, src_lang, tgt_lang, document_id)
        elif file_type == '.xls' or file_type == '.xlsx':
            path, document_save_name = self.file_convert.xls_to_xlsx(Config.DOC_UPLOAD_FOLDER, document_save_name)
            self.excelTrans(path, document_save_name, src_lang, tgt_lang, document_id)
        elif file_type == '.txt':
            isSplit = False
            path = os.path.join(Config.DOC_UPLOAD_FOLDER, document_save_name)
            # path = Config.DOC_UPLOAD_FOLDER + document_save_name
            sub_file_nums = math.ceil(int(document_size) / 1048576)
           
            if sub_file_nums >= 2:
                isSplit = True
                split_file_names = self.utils.splitFile(path, document_id)
                fileThreads = []
                error_files = []
                for split_file_name in split_file_names:
                    txt_file_path = os.path.join(Config.DOC_SPLIT_FOLDER, document_id, split_file_name)
                    try:
                        self.txtTrans(txt_file_path, document_save_name, src_lang, tgt_lang, document_id, isSplit)
                    except Exception as e:
                        print("First translate error: "+ txt_file_path + " " + str(e))
                        error_files.append(txt_file_path)
                
                if len(error_files) != 0:
                    for error_file in error_files:
                        try:
                            print("TRY AGAIN: " + error_file)
                            self.txtTrans(error_file, document_save_name, src_lang, tgt_lang, document_id, isSplit)
                        except Exception as e:
                            print("Second translate error: "+ error_file + " " + str(e))
                            self.p.updateDocumentStatus(document_id, '3')
                    # aThread = threading.Thread(target=self.txtTrans,args=(txt_file_path, document_save_name, src_lang, tgt_lang, document_id, isSplit))
                    # fileThreads.append(aThread)
                
                # event = threading.Event()
            
                # for thread in fileThreads:
                    # thread.start()
                    # event.wait(5)
                
                # for t in fileThreads:
                    # t.join()
                
                try:
                    self.utils.concatFile(os.path.join(Config.DOC_SPLIT_FOLDER, document_id, 'concat'), document_save_name)
                except Exception as e:
                    print(e)
                    self.p.updateDocumentStatus(document_id, '3')
                # self.p.updateDocumentStatus(document_id, '2')
            else:
                try:
                    self.txtTrans(path, document_save_name, src_lang, tgt_lang, document_id, isSplit)
                except Exception as e:
                    print(e)
                    self.p.updateDocumentStatus(document_id, '3')

            self.p.updateDocumentStatus(document_id, '2')

        self.p.updateDocumentSaveName(document_id, document_save_name)
